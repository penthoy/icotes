"""
Word document handler

Supports reading and writing Microsoft Word files (.docx)
"""

import io
import logging
from typing import Any, Dict, List, Optional, Union

from .base import (
    DocumentFormat,
    DocumentHandler,
    HandlerResult,
    register_handler,
)

logger = logging.getLogger(__name__)


class WordHandler(DocumentHandler):
    """Handler for Word documents (.docx)"""
    
    @property
    def supported_formats(self) -> List[DocumentFormat]:
        return [DocumentFormat.WORD_DOCX]
    
    async def read_content(
        self,
        file_data: bytes,
        file_path: str,
        options: Optional[Dict[str, Any]] = None
    ) -> HandlerResult:
        """
        Extract content from Word document.
        
        Options:
            include_headers: Include header text (default: True)
            include_footers: Include footer text (default: False)
            include_tables: Extract tables separately (default: True)
            max_paragraphs: Maximum paragraphs to read (default: None)
        """
        options = options or {}
        include_headers = options.get("include_headers", True)
        include_footers = options.get("include_footers", False)
        include_tables = options.get("include_tables", True)
        max_paragraphs = options.get("max_paragraphs")
        
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            return HandlerResult(
                success=False,
                error="python-docx not installed. Install with: pip install python-docx"
            )
        
        try:
            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)
            
            content_parts = []
            tables = []
            
            # Extract headers if requested
            if include_headers:
                for section in doc.sections:
                    header = section.header
                    if header and header.paragraphs:
                        header_text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
                        if header_text:
                            content_parts.append(f"[HEADER]\n{header_text}\n")
            
            # Extract main body paragraphs
            para_count = 0
            for para in doc.paragraphs:
                if max_paragraphs and para_count >= max_paragraphs:
                    content_parts.append(f"... (truncated at {max_paragraphs} paragraphs)")
                    break
                    
                text = para.text.strip()
                if text:
                    # Add style info for headings
                    if para.style and para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        prefix = '#' * int(level) if level.isdigit() else '##'
                        content_parts.append(f"{prefix} {text}")
                    else:
                        content_parts.append(text)
                    para_count += 1
            
            # Extract tables if requested
            if include_tables:
                for idx, table in enumerate(doc.tables):
                    table_data = self._extract_table(table)
                    if table_data:
                        tables.append({
                            "table_index": idx,
                            "columns": table_data["columns"],
                            "row_count": len(table_data["data"]),
                            "data": table_data["data"]
                        })
                        # Also add table text to content
                        content_parts.append(f"\n[TABLE {idx + 1}]")
                        content_parts.append(self._table_to_text(table_data))
            
            # Extract footers if requested
            if include_footers:
                for section in doc.sections:
                    footer = section.footer
                    if footer and footer.paragraphs:
                        footer_text = "\n".join(p.text for p in footer.paragraphs if p.text.strip())
                        if footer_text:
                            content_parts.append(f"\n[FOOTER]\n{footer_text}")
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Estimate page count (rough approximation)
            # Word typically has ~3000 chars per page
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            estimated_pages = max(1, total_chars // 3000 + 1)
            
            return HandlerResult(
                success=True,
                content="\n\n".join(content_parts),
                metadata=metadata,
                tables=tables,
                pages=estimated_pages
            )
            
        except PackageNotFoundError:
            return HandlerResult(
                success=False,
                error="Invalid or corrupted Word document"
            )
        except Exception as e:
            logger.error(f"Word document read error: {e}")
            return HandlerResult(
                success=False,
                error=f"Failed to read Word document: {str(e)}"
            )
    
    def _extract_table(self, table) -> Optional[Dict[str, Any]]:
        """Extract data from a Word table"""
        try:
            rows_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_data)
            
            if not rows_data:
                return None
            
            # Use first row as headers
            headers = rows_data[0] if rows_data else []
            headers = [h if h else f"col_{i}" for i, h in enumerate(headers)]
            
            data = [
                dict(zip(headers, row))
                for row in rows_data[1:]
            ]
            
            return {"columns": headers, "data": data}
        except Exception as e:
            logger.debug(f"Table extraction error: {e}")
            return None
    
    def _table_to_text(self, table_data: Dict[str, Any]) -> str:
        """Convert table data to readable text format"""
        if not table_data:
            return ""
        
        columns = table_data.get("columns", [])
        data = table_data.get("data", [])
        
        lines = ["\t".join(columns)]
        lines.append("-" * (len(columns) * 15))
        
        for row in data:
            line = "\t".join(str(row.get(col, "")) for col in columns)
            lines.append(line)
        
        return "\n".join(lines)
    
    def _extract_metadata(self, doc) -> Dict[str, Any]:
        """Extract document metadata"""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
            }
        except Exception as e:
            logger.debug(f"Metadata extraction error: {e}")
        
        return metadata
    
    async def write_content(
        self,
        content: Union[str, Dict[str, Any]],
        file_path: str,
        options: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Create Word document from content.
        
        Content can be:
            - Plain text string (creates paragraphs)
            - Dict with structure:
                {
                    "paragraphs": ["text1", "text2", ...],
                    "title": "Document Title",
                    "headings": [{"level": 1, "text": "Heading"}]
                }
        
        Options:
            title: Document title
            author: Document author
        """
        options = options or {}
        
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise ImportError(
                "python-docx required for Word creation. "
                "Install with: pip install python-docx"
            )
        
        doc = Document()
        
        # Set metadata
        doc.core_properties.title = options.get("title", "")
        doc.core_properties.author = options.get("author", "")
        
        if isinstance(content, str):
            # Split text into paragraphs and add them
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    # Check for markdown-style headings
                    if para_text.startswith("# "):
                        doc.add_heading(para_text[2:], level=1)
                    elif para_text.startswith("## "):
                        doc.add_heading(para_text[3:], level=2)
                    elif para_text.startswith("### "):
                        doc.add_heading(para_text[4:], level=3)
                    else:
                        doc.add_paragraph(para_text.strip())
        
        elif isinstance(content, dict):
            # Handle structured content
            if content.get("title"):
                doc.add_heading(content["title"], level=0)
            
            for item in content.get("paragraphs", []):
                if isinstance(item, str):
                    doc.add_paragraph(item)
                elif isinstance(item, dict):
                    if item.get("heading"):
                        doc.add_heading(item["text"], level=item.get("level", 1))
                    else:
                        doc.add_paragraph(item.get("text", ""))
            
            # Add tables if present
            for table_data in content.get("tables", []):
                self._add_table(doc, table_data)
        
        else:
            raise ValueError(f"Unsupported content type: {type(content)}")
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
    
    def _add_table(self, doc, table_data: Dict[str, Any]) -> None:
        """Add a table to the document"""
        columns = table_data.get("columns", [])
        data = table_data.get("data", [])
        
        if not columns or not data:
            return
        
        table = doc.add_table(rows=len(data) + 1, cols=len(columns))
        table.style = 'Table Grid'
        
        # Add header row
        header_row = table.rows[0]
        for idx, col in enumerate(columns):
            header_row.cells[idx].text = str(col)
        
        # Add data rows
        for row_idx, row_data in enumerate(data):
            row = table.rows[row_idx + 1]
            for col_idx, col in enumerate(columns):
                row.cells[col_idx].text = str(row_data.get(col, ""))


# Register handler when module is imported
_handler = WordHandler()
register_handler(_handler)
